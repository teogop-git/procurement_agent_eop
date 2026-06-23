"""
Document Processor - eop.bg downloader + structured document extraction.

This version keeps the old public API (process_email -> combined_text), but also
adds the adapter needed by the ZOP v2 pipeline:

    fetch_procurement_documents(url) -> List[raw_document]

raw_document format:
{
  "procurement_url": "https://app.eop.bg/today/...",
  "tender_id": "...",
  "title": "...",
  "document_name": "technical_specification.pdf",
  "document_type": "technical_specification",
  "url": "https://...signed-url-or-source...",
  "local_path": "/app/output/downloads/.../file.pdf",
  "extraction_quality": "text | empty | unsupported | error",
  "pages": [
    {"page": 1, "section": null, "text": "..."}
  ]
}

The old code returned one giant combined_text block. That was not enough for
traceable ZOP reasoning, because the analyzer could not cite document/page/chunk
sources. This version preserves document/page metadata so downstream analysis can
produce an evidence ledger.
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import requests
import urllib3
from bs4 import BeautifulSoup
from docx import Document
import PyPDF2

try:
    from openpyxl import load_workbook
except Exception:  # openpyxl is optional but recommended for .xlsx extraction
    load_workbook = None

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

logger = logging.getLogger("DocumentProcessor")

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "application/json, text/plain, */*",
    "Origin": "https://app.eop.bg",
    "Referer": "https://app.eop.bg/",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".xlsx"}
ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z"}

MAX_ARCHIVE_DEPTH = 3
MAX_EXTRACTED_FILES = 200
MAX_ARCHIVE_SIZE_BYTES = 500 * 1024 * 1024
MAX_ZIP_UNCOMPRESSED_BYTES = 2 * 1024 * 1024 * 1024


class DocumentProcessor:
    def __init__(self, download_dir: str = "/app/output/downloads"):
        self.download_dir = Path(download_dir)
        self.download_dir.mkdir(parents=True, exist_ok=True)

    # ── Tender ID ─────────────────────────────────────────────────────────────
    @staticmethod
    def _extract_tender_id(url: str) -> Optional[str]:
        for pattern in (
            r"/today/(\d+)",
            r"/(?:procurements|tenders?|procurement)/(\d+)",
            r"/(\d{4,})(?:[/?#]|$)",
        ):
            m = re.search(pattern, url or "")
            if m:
                return m.group(1)
        return None

    @staticmethod
    def _safe_filename(name: str, fallback: str = "document") -> str:
        name = name or fallback
        name = name.split("?")[0].split("#")[0]
        name = re.sub(r"[^\w\-.() А-Яа-я]+", "_", name, flags=re.UNICODE).strip("._ ")
        return name[:180] or fallback

    @staticmethod
    def _guess_document_type(name: str = "", text_sample: str = "") -> str:
        name_l = (name or "").lower()
        text_l = (text_sample or "")[:3000].lower()
        joined = f"{name_l} {text_l}"

        # Filename has higher priority than body text. Price templates often contain
        # generic technical wording, so they must be classified before technical specs.
        if any(k in name_l for k in ["ценово предложение", "ценова оферта", "образец цена", "предлагана цена"]):
            return "price_offer_template"

        if any(k in name_l for k in ["техническа спецификация", "технически спецификации", "техн. спецификация", "technical specification"]):
            return "technical_specification"

        if any(k in name_l for k in ["аргументация", "прогнозна стойност", "стойност"]):
            return "estimated_value_argumentation"

        if any(k in name_l for k in ["методика", "оценка", "комплексна оценка"]):
            return "evaluation_methodology"

        if any(k in name_l for k in ["проект на договор", "договор"]):
            return "contract_draft"

        if any(k in name_l for k in ["декларация", "еедоп"]):
            return "declaration_template"

        if any(k in name_l for k in ["разяснение", "въпрос", "отговор"]):
            return "clarification"

        if any(k in name_l for k in ["покана", "обява", "решение", "обявление", "график", "параметри"]):
            return "invitation"

        # Body fallback.
        rules = [
            ("price_offer_template", ["ценово предложение", "ценова оферта", "образец цена", "предлагана цена"]),
            ("technical_specification", ["техническа спецификация", "технически изисквания", "минимални технически изисквания", "минимални изисквания"]),
            ("estimated_value_argumentation", ["аргументация", "прогнозна стойност"]),
            ("evaluation_methodology", ["методика", "комплексна оценка", "показател за оценка"]),
            ("contract_draft", ["проект на договор", "гаранция за изпълнение"]),
            ("declaration_template", ["декларация", "еедоп", "единен европейски документ"]),
            ("clarification", ["разяснение", "въпрос", "отговор"]),
            ("invitation", ["покана", "обява", "решение", "обявление"]),
        ]

        for doc_type, keys in rules:
            if any(k in joined for k in keys):
                return doc_type

        return "other"

    # ── Playwright download ───────────────────────────────────────────────────
    def _download_all_documents_via_playwright(self, url: str, tender_id: str, dest_dir: Path) -> List[Path]:
        """
        Downloads individual EOP attachment documents, not only the published export ZIP.

        Why this exists:
        app.eop.bg often shows the real tender files in a modal such as
        "Съобщение към поръчката". Those files are separate attachments and are not
        necessarily part of GetPublishedTenderExportsByTenderId ZIP.

        Strategy:
        1. Open the page with a real Chromium browser.
        2. Capture all GetSignedUrlByDocumentId responses.
        3. Recursively scan JSON API responses for objects containing DocumentId + Name.
        4. Open likely document/message modals and click all visible download controls.
        5. For discovered DocumentId values, call GetSignedUrlByDocumentId from inside
           the browser context, then download the signed URLs with requests.
        """
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            logger.warning("Playwright не е инсталиран – не мога да сваля индивидуалните приложения")
            return []

        dest_dir.mkdir(parents=True, exist_ok=True)
        downloaded_paths: List[Path] = []
        signed_items: List[Dict[str, str]] = []
        document_candidates: Dict[str, Dict[str, Any]] = {}

        def _name_looks_like_document(name: str) -> bool:
            name_l = (name or "").lower()
            return any(name_l.endswith(ext) for ext in sorted(SUPPORTED_EXTENSIONS | ARCHIVE_EXTENSIONS | {".xls"}))

        def _add_candidate(obj: Dict[str, Any]) -> None:
            doc_id = obj.get("DocumentId") or obj.get("documentId")
            name = (
                obj.get("Name")
                or obj.get("name")
                or obj.get("FileName")
                or obj.get("fileName")
                or obj.get("DocumentName")
                or obj.get("documentName")
            )
            if not doc_id or not name:
                return
            if not _name_looks_like_document(str(name)):
                return
            document_candidates[str(doc_id)] = {
                "documentId": int(doc_id),
                "name": str(name),
            }

        def _scan_json(obj: Any) -> None:
            if isinstance(obj, dict):
                _add_candidate(obj)
                for value in obj.values():
                    _scan_json(value)
            elif isinstance(obj, list):
                for item in obj:
                    _scan_json(item)

        def _add_signed_item(data: Dict[str, Any], fallback_name: Optional[str] = None) -> None:
            signed_url = (data.get("Url") or data.get("url") or "").replace("\\/", "/")
            name = data.get("Name") or data.get("name") or fallback_name or f"document_{len(signed_items)+1}.bin"
            if signed_url:
                signed_items.append({"name": str(name), "url": signed_url})
                logger.info("  Прихванат signed URL за приложение: %s", name)

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch(
                    headless=True,
                    args=[
                        "--no-sandbox",
                        "--disable-setuid-sandbox",
                        "--disable-dev-shm-usage",
                        "--disable-gpu",
                    ],
                )
                ctx = browser.new_context(
                    user_agent=HEADERS["User-Agent"],
                    ignore_https_errors=True,
                    viewport={"width": 1440, "height": 1000},
                    accept_downloads=True,
                )
                page = ctx.new_page()

                def on_download(download):
                    try:
                        safe = self._safe_filename(download.suggested_filename, fallback=f"download_{len(downloaded_paths)+1}")
                        target = dest_dir / safe
                        download.save_as(str(target))
                        downloaded_paths.append(target)
                        logger.info("  ✅ Browser download: %s", safe)
                    except Exception as exc:
                        logger.debug("  Browser download save failed: %s", exc)

                def on_response(response):
                    try:
                        response_url = response.url
                        if "GetSignedUrlByDocumentId" in response_url:
                            data = response.json()
                            if isinstance(data, dict):
                                _add_signed_item(data)
                            return

                        # Scan only likely EOP API JSON responses. Avoid parsing unrelated assets.
                        if "eop.bg" not in response_url:
                            return
                        content_type = (response.headers.get("content-type") or "").lower()
                        if "json" not in content_type and "NX1Service" not in response_url:
                            return
                        data = response.json()
                        _scan_json(data)
                    except Exception:
                        pass

                page.on("download", on_download)
                page.on("response", on_response)

                logger.info("  Playwright зарежда за индивидуални приложения: %s", url)
                page.goto(url, timeout=60_000, wait_until="networkidle")
                page.wait_for_timeout(3000)

                # Try to open document/message areas. EOP UI changes often; keep this broad but controlled.
                text_patterns = [
                    "Съобщение към поръчката",
                    "Документи",
                    "Приложения",
                    "Приложение",
                    "Покана",
                    "Разяснения",
                    "Обява",
                    "Решение",
                ]
                for pattern in text_patterns:
                    try:
                        loc = page.get_by_text(pattern, exact=False)
                        count = min(loc.count(), 8)
                        for i in range(count):
                            try:
                                loc.nth(i).click(timeout=1500)
                                page.wait_for_timeout(800)
                            except Exception:
                                pass
                    except Exception:
                        pass

                # Click visible download controls in currently opened page/modals.
                click_script = """
                async () => {
                  function visible(el) {
                    const r = el.getBoundingClientRect();
                    const st = window.getComputedStyle(el);
                    return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
                  }
                  function textOf(el) {
                    return ((el.innerText || '') + ' ' + (el.title || '') + ' ' +
                            (el.getAttribute('aria-label') || '') + ' ' +
                            (el.className || '') + ' ' + (el.outerHTML || '').slice(0, 500)).toLowerCase();
                  }
                  const raw = Array.from(document.querySelectorAll(
                    'button, a, [role="button"], [class*="pointer"], [class*="download"], clr-icon, i, svg'
                  ));
                  const controls = [];
                  const seen = new Set();
                  for (const el of raw) {
                    let t = textOf(el);
                    const looksDownload =
                      t.includes('download') || t.includes('изтег') || t.includes('свали') ||
                      t.includes('mdi-download') || t.includes('fa-download') ||
                      t.includes('shape="download"') || t.includes('download"></clr-icon') ||
                      t.includes('icon-download');
                    if (!looksDownload) continue;
                    let target = el.closest('button, a, [role="button"], [class*="pointer"]') || el;
                    if (!visible(target)) continue;
                    if (seen.has(target)) continue;
                    seen.add(target);
                    controls.push(target);
                  }
                  let clicked = 0;
                  for (const el of controls.slice(0, 80)) {
                    try {
                      el.scrollIntoView({block: 'center', inline: 'center'});
                      await new Promise(r => setTimeout(r, 150));
                      el.click();
                      clicked += 1;
                      await new Promise(r => setTimeout(r, 400));
                    } catch (e) {}
                  }
                  return clicked;
                }
                """
                try:
                    clicked = page.evaluate(click_script)
                    logger.info("  Playwright: натиснати download контроли: %s", clicked)
                    page.wait_for_timeout(5000)
                except Exception as exc:
                    logger.debug("  Click download controls failed: %s", exc)

                # Use discovered DocumentId values to request signed URLs inside browser context.
                if document_candidates:
                    logger.info("  Намерени DocumentId кандидати: %s", len(document_candidates))
                    docs_payload = list(document_candidates.values())[:120]
                    try:
                        signed_from_candidates = page.evaluate(
                            """
                            async (docs) => {
                              const out = [];
                              for (const d of docs) {
                                try {
                                  const response = await fetch('https://service.eop.bg/NX1Service.svc/GetSignedUrlByDocumentId', {
                                    method: 'POST',
                                    credentials: 'include',
                                    headers: {
                                      'Content-Type': 'application/json; charset=UTF-8',
                                      'X-Requested-With': 'XMLHttpRequest'
                                    },
                                    body: JSON.stringify({documentId: d.documentId})
                                  });
                                  const txt = await response.text();
                                  let data = {};
                                  try { data = JSON.parse(txt); } catch (e) { data = {}; }
                                  if (data.Url || data.url) {
                                    out.push({
                                      name: data.Name || data.name || d.name,
                                      url: String(data.Url || data.url).replaceAll('\\\\/', '/')
                                    });
                                  }
                                } catch (e) {}
                              }
                              return out;
                            }
                            """,
                            docs_payload,
                        )
                        for item in signed_from_candidates or []:
                            if item.get("url"):
                                signed_items.append({"name": item.get("name") or "document.bin", "url": item["url"]})
                    except Exception as exc:
                        logger.debug("  Browser-context GetSignedUrl failed: %s", exc)

                browser.close()
        except Exception as exc:
            logger.error("  Playwright индивидуални приложения грешка: %s", exc)
            return downloaded_paths

        # Download signed URLs; de-duplicate by URL and filename.
        seen_urls = set()
        seen_names = {p.name for p in downloaded_paths}
        for item in signed_items:
            signed_url = item.get("url") or ""
            name = item.get("name") or "document.bin"
            if not signed_url or signed_url in seen_urls:
                continue
            seen_urls.add(signed_url)

            safe_name = self._safe_filename(name, fallback=f"document_{len(downloaded_paths)+1}.bin")
            if safe_name in seen_names:
                stem = Path(safe_name).stem
                suffix = Path(safe_name).suffix
                safe_name = f"{stem}_{len(seen_names)+1}{suffix}"
            seen_names.add(safe_name)

            downloaded = self._download_from_url(signed_url, safe_name, dest_dir)
            if downloaded:
                downloaded_paths.append(downloaded)

        logger.info("  ✅ Индивидуално свалени приложения: %s", len(downloaded_paths))
        return downloaded_paths

    def _download_via_playwright(self, url: str, tender_id: str, dest_dir: Path) -> Optional[Path]:
        """Use Playwright to capture the eop.bg signed URL and download the exported ZIP."""
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            logger.warning("Playwright не е инсталиран – пропускам Playwright fallback")
            return None

        signed_url = None
        zip_name = f"T{tender_id}-export.zip"

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch(
                    headless=True,
                    args=[
                        "--no-sandbox",
                        "--disable-setuid-sandbox",
                        "--disable-dev-shm-usage",
                        "--disable-gpu",
                    ],
                )
                ctx = browser.new_context(
                    user_agent=HEADERS["User-Agent"],
                    ignore_https_errors=True,
                    viewport={"width": 1440, "height": 1000},
                )
                page = ctx.new_page()

                def on_response(response):
                    nonlocal signed_url, zip_name
                    if "GetSignedUrlByDocumentId" not in response.url:
                        return
                    try:
                        data = response.json()
                        if data.get("Url"):
                            signed_url = data["Url"].replace("\\/", "/")
                            zip_name = data.get("Name", zip_name)
                            logger.info("  Прихванат signed URL за: %s", zip_name)
                    except Exception as exc:
                        logger.debug("  Неуспешно четене на GetSignedUrl response: %s", exc)

                page.on("response", on_response)

                logger.info("  Playwright зарежда: %s", url)
                page.goto(url, timeout=45_000, wait_until="networkidle")
                page.wait_for_timeout(1500)

                if not signed_url:
                    click_selectors = [
                        "text=.zip",
                        "a:has-text('.zip')",
                        "[class*='clarityblue'][class*='pointer']",
                        "button:has-text('Изтегли')",
                    ]
                    for selector in click_selectors:
                        try:
                            locator = page.locator(selector).first
                            if locator.count() == 0:
                                continue
                            with page.expect_response(
                                lambda r: "GetSignedUrlByDocumentId" in r.url,
                                timeout=12_000,
                            ):
                                locator.click()
                            page.wait_for_timeout(2000)
                            if signed_url:
                                break
                        except Exception as exc:
                            logger.debug("  Playwright click selector failed [%s]: %s", selector, exc)

                browser.close()
        except Exception as exc:
            logger.error("  Playwright грешка: %s", exc)
            return None

        if not signed_url:
            logger.warning("  Playwright: signed URL не е намерен")
            return None

        return self._download_from_url(signed_url, zip_name, dest_dir)

    # ── Direct service.eop.bg ────────────────────────────────────────────────
    def _download_via_service_api(self, tender_id: str, dest_dir: Path) -> Optional[Path]:
        """
        Direct eop.bg API path. Works only when the caller IP is allowed by service.eop.bg.
        If blocked, caller will fall back to Playwright.
        """
        headers = {
            **HEADERS,
            "Content-Type": "application/json; charset=UTF-8",
            "X-Requested-With": "XMLHttpRequest",
            "sec-ch-ua": '"Google Chrome";v="147", "Not.A/Brand";v="8"',
            "Sec-Fetch-Site": "same-site",
            "Sec-Fetch-Mode": "cors",
            "Sec-Fetch-Dest": "empty",
        }

        try:
            r1 = requests.post(
                "https://service.eop.bg/NX1Service.svc/GetPublishedTenderExportsByTenderId",
                headers=headers,
                json={"tenderId": int(tender_id), "ianaTimeZone": "Europe/Sofia"},
                timeout=20,
                verify=False,
            )
            if r1.status_code != 200 or not r1.text.lstrip().startswith("["):
                logger.warning("  API стъпка 1: %s / %s", r1.status_code, r1.text[:120])
                return None

            exports = r1.json()
            if not exports:
                logger.warning("  API: няма exports за тази поръчка")
                return None

            export = exports[0]
            doc_id = export["DocumentId"]
            zip_name = export.get("Name", f"T{tender_id}-export.zip")
            logger.info("  DocumentId: %s / %s", doc_id, zip_name)
        except Exception as exc:
            logger.warning("  API стъпка 1 грешка: %s", exc)
            return None

        try:
            r2 = requests.post(
                "https://service.eop.bg/NX1Service.svc/GetSignedUrlByDocumentId",
                headers=headers,
                json={"documentId": doc_id},
                timeout=20,
                verify=False,
            )
            if r2.status_code != 200:
                logger.warning("  API стъпка 2: %s / %s", r2.status_code, r2.text[:120])
                return None

            data = r2.json()
            signed_url = data.get("Url", "").replace("\\/", "/")
            if not signed_url:
                logger.warning("  API: няма Url в response")
                return None

            logger.info("  Signed URL получен")
            return self._download_from_url(signed_url, zip_name, dest_dir)
        except Exception as exc:
            logger.warning("  API стъпка 2 грешка: %s", exc)
            return None

    # ── HTTP download ─────────────────────────────────────────────────────────
    def _download_from_url(self, url: str, filename: str, dest_dir: Path) -> Optional[Path]:
        try:
            dest_dir.mkdir(parents=True, exist_ok=True)
            safe_name = self._safe_filename(filename, fallback="eop_export.zip")
            dest = dest_dir / safe_name

            with requests.get(url, timeout=180, stream=True, verify=False) as resp:
                resp.raise_for_status()
                with open(dest, "wb") as f:
                    for chunk in resp.iter_content(1024 * 1024):
                        if chunk:
                            f.write(chunk)

            size_kb = dest.stat().st_size // 1024
            logger.info("  ✅ Свален: %s (%s KB)", safe_name, size_kb)
            return dest
        except Exception as exc:
            logger.error("  Download грешка: %s", exc)
            return None

    # ── Text extraction ───────────────────────────────────────────────────────
    def _extract_pdf_pages(self, path: Path) -> List[Dict[str, Any]]:
        pages: List[Dict[str, Any]] = []
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for idx, page in enumerate(reader.pages, start=1):
                    text = page.extract_text() or ""
                    pages.append({"page": idx, "section": None, "text": text.strip()})
        except Exception as exc:
            logger.error("  PDF грешка %s: %s", path.name, exc)
        return pages

    def _extract_docx_pages(self, path: Path) -> List[Dict[str, Any]]:
        try:
            doc = Document(str(path))
            parts: List[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    parts.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            text = "\n".join(parts).strip()
            return [{"page": 1, "section": None, "text": text}] if text else []
        except Exception as exc:
            logger.error("  DOCX грешка %s: %s", path.name, exc)
            return []

    def _extract_doc_legacy_pages(self, path: Path) -> List[Dict[str, Any]]:
        # Attempt 1: antiword
        if shutil.which("antiword"):
            try:
                result = subprocess.run(
                    ["antiword", str(path)],
                    capture_output=True,
                    text=True,
                    timeout=45,
                    encoding="utf-8",
                    errors="replace",
                )
                if result.returncode == 0 and result.stdout.strip():
                    logger.info("  .doc чрез antiword: %s", path.name)
                    return [{"page": 1, "section": None, "text": result.stdout.strip()}]
            except Exception as exc:
                logger.warning("  antiword грешка: %s", exc)

        # Attempt 2: LibreOffice -> docx
        libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
        if libreoffice:
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    subprocess.run(
                        [
                            libreoffice,
                            "--headless",
                            "--convert-to",
                            "docx",
                            "--outdir",
                            tmpdir,
                            str(path),
                        ],
                        capture_output=True,
                        timeout=90,
                    )
                    converted = Path(tmpdir) / f"{path.stem}.docx"
                    if converted.exists():
                        pages = self._extract_docx_pages(converted)
                        if pages:
                            logger.info("  .doc чрез LibreOffice: %s", path.name)
                            return pages
            except Exception as exc:
                logger.warning("  LibreOffice грешка: %s", exc)

        logger.warning("  Не може да се прочете .doc: %s", path.name)
        return []

    def _extract_text_pages(self, path: Path) -> List[Dict[str, Any]]:
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            text = raw.strip()
            return [{"page": 1, "section": None, "text": text}] if text else []
        except Exception as exc:
            logger.warning("  TXT/RTF грешка %s: %s", path.name, exc)
            return []

    def _extract_xlsx_pages(self, path: Path) -> List[Dict[str, Any]]:
        if load_workbook is None:
            logger.warning("  openpyxl не е наличен – пропускам XLSX: %s", path.name)
            return []

        try:
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)
            pages: List[Dict[str, Any]] = []

            for idx, ws in enumerate(wb.worksheets, start=1):
                lines: List[str] = [f"=== Sheet: {ws.title} ==="]
                for row in ws.iter_rows(values_only=True):
                    values = []
                    for value in row:
                        if value is None:
                            continue
                        value_s = str(value).strip()
                        if value_s:
                            values.append(value_s)
                    if values:
                        lines.append(" | ".join(values))

                text = "\n".join(lines).strip()
                if text:
                    pages.append({"page": idx, "section": ws.title, "text": text})

            return pages
        except Exception as exc:
            logger.warning("  XLSX грешка %s: %s", path.name, exc)
            return []

    def _extract_pages_for_file(self, path: Path) -> List[Dict[str, Any]]:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf_pages(path)
        if ext == ".docx":
            return self._extract_docx_pages(path)
        if ext == ".doc":
            return self._extract_doc_legacy_pages(path)
        if ext in {".txt", ".rtf"}:
            return self._extract_text_pages(path)
        if ext == ".xlsx":
            return self._extract_xlsx_pages(path)
        return []

    @staticmethod
    def _is_relative_to(child: Path, parent: Path) -> bool:
        try:
            child.resolve().relative_to(parent.resolve())
            return True
        except ValueError:
            return False

    @staticmethod
    def _safe_extract_dir(dest_dir: Path, archive_path: Path) -> Path:
        base = dest_dir / f"{archive_path.stem}_extracted"
        if not base.exists():
            base.mkdir(parents=True, exist_ok=True)
            return base

        for idx in range(2, 1000):
            candidate = dest_dir / f"{archive_path.stem}_extracted_{idx}"
            if not candidate.exists():
                candidate.mkdir(parents=True, exist_ok=True)
                return candidate

        raise RuntimeError(f"Не мога да създам extract директория за {archive_path.name}")

    def _extract_zip_to_dir(self, archive_path: Path, extract_dir: Path) -> None:
        with zipfile.ZipFile(archive_path, "r") as zf:
            members = zf.infolist()
            logger.info("  ZIP: %s файла", len(members))

            if len(members) > MAX_EXTRACTED_FILES:
                raise RuntimeError(
                    f"ZIP архивът съдържа твърде много файлове: {len(members)} > {MAX_EXTRACTED_FILES}"
                )

            total_uncompressed = sum(member.file_size for member in members)
            if total_uncompressed > MAX_ZIP_UNCOMPRESSED_BYTES:
                raise RuntimeError(
                    "ZIP архивът е твърде голям след разархивиране: "
                    f"{total_uncompressed} bytes > {MAX_ZIP_UNCOMPRESSED_BYTES} bytes"
                )

            extract_root = extract_dir.resolve()
            for member in members:
                member_name = member.filename
                if not member_name or member_name.endswith("/"):
                    continue

                target = (extract_dir / member_name).resolve()
                if not self._is_relative_to(target, extract_root):
                    raise RuntimeError(f"Небезопасен ZIP member path: {member_name}")

                zf.extract(member, extract_dir)

    def _archive_tool_commands(self, archive_path: Path, extract_dir: Path) -> List[List[str]]:
        commands: List[List[str]] = []

        seven_zip = shutil.which("7z") or shutil.which("7za")
        if seven_zip:
            commands.append([
                seven_zip,
                "x",
                "-y",
                "-bd",
                f"-o{str(extract_dir)}",
                str(archive_path),
            ])

        unar = shutil.which("unar")
        if unar:
            commands.append([
                unar,
                "-quiet",
                "-force-overwrite",
                "-output-directory",
                str(extract_dir),
                str(archive_path),
            ])

        unrar = shutil.which("unrar")
        if unrar:
            commands.append([
                unrar,
                "x",
                "-o+",
                "-y",
                str(archive_path),
                str(extract_dir) + "/",
            ])

        return commands

    def _extract_external_archive_to_dir(self, archive_path: Path, extract_dir: Path) -> None:
        commands = self._archive_tool_commands(archive_path, extract_dir)
        if not commands:
            raise RuntimeError(
                "Няма наличен инструмент за разархивиране на RAR/7Z. "
                "Инсталирай 7z/7za, unar или unrar в контейнера."
            )

        last_error = ""
        for cmd in commands:
            try:
                logger.info("  Архив чрез %s: %s", Path(cmd[0]).name, archive_path.name)
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=180,
                    encoding="utf-8",
                    errors="replace",
                )
                if result.returncode == 0:
                    return

                last_error = (
                    f"cmd={' '.join(cmd)} rc={result.returncode} "
                    f"stdout={result.stdout[-1500:]} stderr={result.stderr[-1500:]}"
                )
                logger.warning("  Архив tool failed: %s", last_error)
            except Exception as exc:
                last_error = f"cmd={' '.join(cmd)} exc={exc}"
                logger.warning("  Архив tool exception: %s", last_error)

        raise RuntimeError(f"Неуспешно разархивиране на {archive_path.name}: {last_error}")

    def _extract_archive_to_dir(self, archive_path: Path, extract_dir: Path) -> None:
        ext = archive_path.suffix.lower()

        if archive_path.stat().st_size > MAX_ARCHIVE_SIZE_BYTES:
            raise RuntimeError(
                f"Архивът е твърде голям: {archive_path.stat().st_size} bytes > {MAX_ARCHIVE_SIZE_BYTES} bytes"
            )

        if ext == ".zip":
            self._extract_zip_to_dir(archive_path, extract_dir)
            return

        if ext in {".rar", ".7z"}:
            self._extract_external_archive_to_dir(archive_path, extract_dir)
            return

        raise RuntimeError(f"Неподдържан архивен формат: {archive_path.name}")

    def _collect_documents_from_extracted_dir(
        self,
        extract_dir: Path,
        source_archive: Path,
        dest_dir: Path,
        procurement_url: str,
        tender_id: Optional[str],
        title: str,
        depth: int,
    ) -> List[Dict[str, Any]]:
        documents: List[Dict[str, Any]] = []
        files_seen = 0

        extract_root = extract_dir.resolve()
        for file_path in sorted(extract_dir.rglob("*")):
            if not file_path.is_file():
                continue

            resolved = file_path.resolve()
            if not self._is_relative_to(resolved, extract_root):
                logger.warning("  Пропускам файл извън extract директорията: %s", file_path)
                continue

            files_seen += 1
            if files_seen > MAX_EXTRACTED_FILES:
                logger.warning(
                    "  Прекратявам обработката на архива %s: повече от %s файла",
                    source_archive.name,
                    MAX_EXTRACTED_FILES,
                )
                break

            ext = file_path.suffix.lower()

            if ext in ARCHIVE_EXTENSIONS:
                if depth >= MAX_ARCHIVE_DEPTH:
                    logger.warning(
                        "  Пропускам nested archive поради MAX_ARCHIVE_DEPTH=%s: %s",
                        MAX_ARCHIVE_DEPTH,
                        file_path.name,
                    )
                    continue

                documents.extend(
                    self._extract_archive_documents(
                        path=file_path,
                        dest_dir=dest_dir,
                        procurement_url=procurement_url,
                        tender_id=tender_id,
                        title=title,
                        depth=depth + 1,
                    )
                )
                continue

            if ext not in SUPPORTED_EXTENSIONS:
                logger.debug("  Пропускам неподдържан файл в архив: %s", file_path.name)
                continue

            doc = self._build_raw_document(
                file_path=file_path,
                procurement_url=procurement_url,
                tender_id=tender_id,
                title=title,
                source_url=str(source_archive),
            )
            documents.append(doc)

        return documents

    def _extract_archive_documents(
        self,
        path: Path,
        dest_dir: Path,
        procurement_url: str,
        tender_id: Optional[str],
        title: str,
        depth: int = 0,
    ) -> List[Dict[str, Any]]:
        documents: List[Dict[str, Any]] = []
        ext = path.suffix.lower()
        archive_kind = ext.lstrip(".").upper() or "ARCHIVE"

        if ext not in ARCHIVE_EXTENSIONS:
            logger.warning("  Неподдържан архивен формат: %s", path.name)
            return documents

        if depth > MAX_ARCHIVE_DEPTH:
            logger.warning("  Пропускам archive depth=%s: %s", depth, path.name)
            return documents

        try:
            extract_dir = self._safe_extract_dir(dest_dir, path)
            self._extract_archive_to_dir(path, extract_dir)

            documents.extend(
                self._collect_documents_from_extracted_dir(
                    extract_dir=extract_dir,
                    source_archive=path,
                    dest_dir=dest_dir,
                    procurement_url=procurement_url,
                    tender_id=tender_id,
                    title=title,
                    depth=depth,
                )
            )
            logger.info("  %s: извлечени документи: %s", archive_kind, len(documents))
        except zipfile.BadZipFile:
            logger.error("  ZIP грешка %s: файлът не е валиден ZIP", path.name)
        except Exception as exc:
            logger.error("  %s грешка %s: %s", archive_kind, path.name, exc)

        return documents

    def _build_raw_document(
        self,
        file_path: Path,
        procurement_url: str,
        tender_id: Optional[str],
        title: str,
        source_url: Optional[str] = None,
    ) -> Dict[str, Any]:
        pages = self._extract_pages_for_file(file_path)
        sample = "\n".join((p.get("text") or "") for p in pages[:2])
        doc_type = self._guess_document_type(file_path.name, sample)

        if not pages:
            extraction_quality = "empty"
        elif len(sample.strip()) < 80:
            extraction_quality = "low_text"
        else:
            extraction_quality = "text"

        logger.info(
            "  Текст от: %s (%s страници, quality=%s, type=%s)",
            file_path.name,
            len(pages),
            extraction_quality,
            doc_type,
        )

        return {
            "procurement_url": procurement_url,
            "tender_id": tender_id,
            "title": title,
            "document_name": file_path.name,
            "document_type": doc_type,
            "url": source_url or procurement_url,
            "local_path": str(file_path),
            "extraction_quality": extraction_quality,
            "pages": pages,
        }

    def _extract_documents(
        self,
        downloaded: Path,
        dest_dir: Path,
        procurement_url: str,
        tender_id: Optional[str],
        title: str,
    ) -> List[Dict[str, Any]]:
        ext = downloaded.suffix.lower()
        if ext in ARCHIVE_EXTENSIONS:
            return self._extract_archive_documents(downloaded, dest_dir, procurement_url, tender_id, title)
        if ext in SUPPORTED_EXTENSIONS:
            return [
                self._build_raw_document(
                    file_path=downloaded,
                    procurement_url=procurement_url,
                    tender_id=tender_id,
                    title=title,
                    source_url=procurement_url,
                )
            ]

        logger.warning("  Неподдържан файл: %s", downloaded.name)
        return [
            {
                "procurement_url": procurement_url,
                "tender_id": tender_id,
                "title": title,
                "document_name": downloaded.name,
                "document_type": "unsupported",
                "url": procurement_url,
                "local_path": str(downloaded),
                "extraction_quality": "unsupported",
                "pages": [],
            }
        ]

    @staticmethod
    def _combine_documents_text(raw_documents: List[Dict[str, Any]]) -> str:
        blocks: List[str] = []
        for doc in raw_documents:
            blocks.append(
                f"=== DOCUMENT: {doc.get('document_name')} | type={doc.get('document_type')} | "
                f"quality={doc.get('extraction_quality')} ==="
            )
            for page in doc.get("pages", []):
                text = (page.get("text") or "").strip()
                if not text:
                    continue
                blocks.append(f"--- PAGE {page.get('page')} ---\n{text}")
        return "\n\n".join(blocks)


    @staticmethod
    def _fix_mojibake(text: str) -> str:
        if not text:
            return text

        value = str(text).strip()

        # Common EOP mojibake: UTF-8 bytes interpreted as latin1/cp1252.
        if "Ð" in value or "Ñ" in value:
            for encoding in ("latin1", "cp1252"):
                try:
                    fixed = value.encode(encoding, errors="strict").decode("utf-8", errors="strict")
                    if fixed and fixed != value:
                        return fixed
                except Exception:
                    pass

        return value

    @staticmethod
    def _normalize_title(title: str, url: str) -> str:
        value = DocumentProcessor._fix_mojibake(title or "").strip()

        # Static EOP HTML often returns only the portal title, not the real tender title.
        generic_titles = {
            "ЦАИС ЕОП",
            "ЕОП",
            "Електронни обществени поръчки",
            "Electronic Public Procurement",
        }

        if not value or value in generic_titles:
            tender_id = DocumentProcessor._extract_tender_id(url)
            return f"ЦАИС ЕОП поръчка {tender_id}" if tender_id else url

        return value[:250]

    def _get_title(self, url: str) -> str:
        try:
            resp = requests.get(url, headers=HEADERS, timeout=20, verify=False)
            soup = BeautifulSoup(resp.text, "html.parser")
            selectors = [
                ".published-tenders-content__title",
                "[class*='published-tenders-content__title']",
                "h1",
                "h2",
                "title",
            ]
            for selector in selectors:
                el = soup.select_one(selector)
                if el:
                    text = el.get_text(" ", strip=True)
                    if text:
                        return self._normalize_title(text, url)
        except Exception as exc:
            logger.debug("  Title scrape грешка: %s", exc)
        return url

    # ── Public API: pipeline adapter ──────────────────────────────────────────
    def fetch_procurement_documents(self, url: str) -> List[Dict[str, Any]]:
        """
        Main adapter for the ZOP v2 pipeline.

        Use this in zop_v2.pipeline.fetch_procurement_documents(url):
            from document_processor import DocumentProcessor
            return DocumentProcessor().fetch_procurement_documents(url)
        """
        logger.info("Обработка: %s", url)
        tender_id = self._extract_tender_id(url)
        logger.info("  Tender ID: %s", tender_id)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = re.sub(r"[^\w]", "_", (tender_id or url)[-60:])
        proc_dir = self.download_dir / ts / safe
        proc_dir.mkdir(parents=True, exist_ok=True)

        title = self._get_title(url)

        downloaded_files: List[Path] = []

        # First try the real page attachments. In many EOP procedures the technical
        # specification is an individual message attachment, not part of the export ZIP.
        logger.info("  → Playwright individual attachments")
        downloaded_files.extend(
            self._download_all_documents_via_playwright(url, tender_id or "0", proc_dir)
        )

        # Fallback to the published tender export ZIP if individual attachments were not found.
        if not downloaded_files and tender_id:
            downloaded = self._download_via_service_api(tender_id, proc_dir)
            if downloaded:
                downloaded_files.append(downloaded)

        if not downloaded_files:
            logger.info("  → Playwright ZIP fallback")
            downloaded = self._download_via_playwright(url, tender_id or "0", proc_dir)
            if downloaded:
                downloaded_files.append(downloaded)

        if not downloaded_files:
            logger.warning("  ⚠ Не е намерен файл за: %s", url)
            return []

        raw_documents: List[Dict[str, Any]] = []
        for downloaded in downloaded_files:
            raw_documents.extend(self._extract_documents(downloaded, proc_dir, url, tender_id, title))

        # De-duplicate extracted docs by filename + first page text prefix.
        deduped: List[Dict[str, Any]] = []
        seen = set()
        for doc in raw_documents:
            first_text = ""
            if doc.get("pages"):
                first_text = (doc["pages"][0].get("text") or "")[:200]
            key = (doc.get("document_name"), first_text)
            if key in seen:
                continue
            seen.add(key)
            if doc.get("pages") or doc.get("extraction_quality") in {"empty", "low_text"}:
                deduped.append(doc)

        logger.info("  ✅ Извлечени документи: %s", len(deduped))
        return deduped

    # ── Public API: backward compatibility ────────────────────────────────────
    def process_url(self, url: str) -> Optional[Dict[str, Any]]:
        raw_documents = self.fetch_procurement_documents(url)
        if not raw_documents:
            return None

        title = raw_documents[0].get("title") or self._get_title(url)
        combined_text = self._combine_documents_text(raw_documents)

        return {
            "procurement_url": url,
            "title": title,
            "combined_text": combined_text,
            "doc_count": len(raw_documents),
            "raw_documents": raw_documents,
            "sources_reviewed": [
                {
                    "document_name": d.get("document_name"),
                    "document_type": d.get("document_type"),
                    "url": d.get("url"),
                    "pages_reviewed": [p.get("page") for p in d.get("pages", []) if p.get("page") is not None],
                    "status": "reviewed" if d.get("pages") else "no_text",
                    "extraction_quality": d.get("extraction_quality"),
                }
                for d in raw_documents
            ],
        }

    def process_email(self, email_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        documents: List[Dict[str, Any]] = []
        for url in email_data.get("procurement_urls", []):
            result = self.process_url(url)
            if result:
                documents.append(result)
                logger.info("  ✅ Обработен: %s", str(result.get("title", ""))[:80])
            else:
                logger.warning("  ⚠ Няма обработваеми документи за: %s", url)
        return documents
