from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Dict

import yaml
from docx import Document

try:
    from procurement_agent.modules.checklist_schema import ChecklistData
except ImportError:
    from checklist_schema import ChecklistData


def normalize_label(value: str) -> str:
    value = value or ""
    value = value.replace("\xa0", " ").replace("\n", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip().strip(":").strip().lower()


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = ""


def set_cell_text(cell, text: str) -> None:
    clear_cell(cell)
    if not cell.paragraphs:
        cell.add_paragraph()
    cell.paragraphs[0].text = text or ""


def iter_table_rows(document: Document):
    for table in document.tables:
        for row in table.rows:
            yield row


def load_mapping(path: Path) -> Dict[str, str]:
    with path.open("r", encoding="utf-8") as f:
        raw = yaml.safe_load(f) or {}

    mapping = raw.get("field_mapping", raw)
    if not isinstance(mapping, dict):
        raise ValueError("Invalid mapping format. Expected dictionary.")

    return {normalize_label(k): str(v) for k, v in mapping.items()}


def populate_checklist(
    template_docx: Path,
    output_docx: Path,
    mapping_yml: Path,
    data_json: Path,
    confidence_threshold: float = 0.85,
    requirement_col: int = 1,
    status_col: int = 3,
) -> Dict[str, Any]:
    document = Document(str(template_docx))
    mapping = load_mapping(mapping_yml)

    with data_json.open("r", encoding="utf-8") as f:
        payload = json.load(f)

    checklist_data = ChecklistData.from_dict(payload)

    report = {
        "template": str(template_docx),
        "output": str(output_docx),
        "updated_rows": [],
        "missing_rows_in_template": [],
        "missing_values": [],
    }

    matched_labels = set()

    for row in iter_table_rows(document):
        if not row.cells:
            continue

        label = normalize_label(row.cells[0].text)
        if not label or label not in mapping:
            continue

        field_key = mapping[label]
        value = checklist_data.get_value(field_key)
        checked_label = checklist_data.get_checked_label(field_key, confidence_threshold)
        matched_labels.add(label)

        if value:
            if len(row.cells) > requirement_col:
                set_cell_text(row.cells[requirement_col], value)
            if len(row.cells) > status_col and checked_label:
                set_cell_text(row.cells[status_col], checked_label)

            report["updated_rows"].append({
                "criteria": row.cells[0].text.strip(),
                "field_key": field_key,
                "status": checked_label,
            })
        else:
            report["missing_values"].append({
                "criteria": row.cells[0].text.strip(),
                "field_key": field_key,
            })

    for label, field_key in mapping.items():
        if label not in matched_labels:
            report["missing_rows_in_template"].append({
                "criteria_normalized": label,
                "field_key": field_key,
            })

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))

    report_path = output_docx.with_suffix(".report.json")
    with report_path.open("w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Populate ZOP checklist DOCX.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--mapping", required=True, type=Path)
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--confidence-threshold", default=0.85, type=float)
    args = parser.parse_args()

    report = populate_checklist(
        template_docx=args.template,
        output_docx=args.output,
        mapping_yml=args.mapping,
        data_json=args.data,
        confidence_threshold=args.confidence_threshold,
    )

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
