"""
Report Generator
Генерира структурирани Word (.docx) репорти от анализа
"""

import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
import docx.oxml as oxml

logger = logging.getLogger("ReportGenerator")

RISK_COLORS = {
    "HIGH":    RGBColor(0xC0, 0x00, 0x00),   # тъмно червено
    "MEDIUM":  RGBColor(0xFF, 0x80, 0x00),   # оранжево
    "LOW":     RGBColor(0x00, 0x70, 0xC0),   # синьо
    "UNKNOWN": RGBColor(0x80, 0x80, 0x80),   # сиво
}


def _set_cell_bg(cell, hex_color: str):
    """Задава фон на клетка от таблица."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = oxml.OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


class ReportGenerator:
    def __init__(self, output_dir: str = "/app/output/reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _heading(self, doc: Document, text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x35, 0x64) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
        return p

    def _risk_paragraph(self, doc: Document, risk: str):
        p = doc.add_paragraph()
        run = p.add_run(f"Ниво на риск: {risk}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RISK_COLORS.get(risk, RISK_COLORS["UNKNOWN"])
        return p

    # ── Sections ──────────────────────────────────────────────────────────────
    def _add_title_page(self, doc: Document, email_data: Dict, results: List[Dict]):
        t = doc.add_heading("VENDOR-LOCK АНАЛИЗ", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        s = doc.add_heading("Обществени Поръчки – Технически Анализ", 2)
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        meta = doc.add_table(rows=4, cols=2)
        meta.style = "Light Grid Accent 1"
        now = datetime.now()
        rows_data = [
            ("Дата на анализ:", now.strftime("%d.%m.%Y %H:%M")),
            ("Имейл тема:", email_data.get("subject", "N/A")),
            ("Брой поръчки:", str(len(results))),
            ("Изпратен от:", email_data.get("sender", "N/A")),
        ]
        for i, (k, v) in enumerate(rows_data):
            meta.cell(i, 0).text = k
            meta.cell(i, 1).text = v
        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, results: List[Dict]):
        self._heading(doc, "EXECUTIVE SUMMARY")

        total = len(results)
        high = sum(1 for r in results if r.get("analysis", {}).get("summary", {}).get("overall_risk_level") == "HIGH")
        medium = sum(1 for r in results if r.get("analysis", {}).get("summary", {}).get("overall_risk_level") == "MEDIUM")

        tbl = doc.add_table(rows=3, cols=2)
        tbl.style = "Light Grid Accent 1"
        tbl.cell(0, 0).text = "Общ брой анализирани поръчки:"
        tbl.cell(0, 1).text = str(total)
        tbl.cell(1, 0).text = "Висок риск (HIGH):"
        tbl.cell(1, 1).text = str(high)
        _set_cell_bg(tbl.cell(1, 1), "FFC7CE")
        tbl.cell(2, 0).text = "Среден риск (MEDIUM):"
        tbl.cell(2, 1).text = str(medium)
        _set_cell_bg(tbl.cell(2, 1), "FFEB9C")
        doc.add_paragraph()

    def _add_procurement_section(self, doc: Document, result: Dict, idx: int):
        title = result.get("title", f"Поръчка {idx + 1}")
        analysis = result.get("analysis", {})
        summary = analysis.get("summary", {})
        risk = summary.get("overall_risk_level", "UNKNOWN")

        self._heading(doc, f"{idx + 1}. {title}", level=1)
        doc.add_paragraph(f"URL: {result.get('procurement_url', 'N/A')}")
        self._risk_paragraph(doc, risk)
        doc.add_paragraph(summary.get("brief_summary", ""))

        # Хардуерни спецификации
        hw = analysis.get("hardware_specifications", [])
        if hw:
            self._heading(doc, "Хардуерни спецификации – Vendor-Lock", 2)
            for item in hw:
                doc.add_paragraph(f"• Изискване: {item.get('requirement', '')}", style="List Bullet")
                doc.add_paragraph(f"  Индикатор: {item.get('vendor_lock_indicator', '')}")
                vendors = item.get("affected_vendors", [])
                if vendors:
                    doc.add_paragraph(f"  Засегнати вендори: {', '.join(vendors)}")
                rec = item.get("recommendation", "")
                if rec:
                    doc.add_paragraph(f"  Препоръка: {rec}")
                doc.add_paragraph()

        # Софтуерни изисквания
        sw = analysis.get("software_requirements", [])
        if sw:
            self._heading(doc, "Софтуерни изисквания – Vendor-Lock", 2)
            for item in sw:
                doc.add_paragraph(f"• {item.get('requirement', '')}", style="List Bullet")
                doc.add_paragraph(f"  Индикатор: {item.get('vendor_lock_indicator', '')}")
                if item.get("recommendation"):
                    doc.add_paragraph(f"  Препоръка: {item['recommendation']}")
                doc.add_paragraph()

        # Сертификации
        certs = analysis.get("certifications_and_standards", [])
        if certs:
            self._heading(doc, "Сертификации и стандарти", 2)
            for item in certs:
                legit = "✓ Легитимно" if item.get("is_legitimate") else "⚠ Проблемно"
                doc.add_paragraph(f"• [{legit}] {item.get('requirement', '')}", style="List Bullet")
                doc.add_paragraph(f"  {item.get('vendor_lock_indicator', '')}")
            doc.add_paragraph()

        # Индиректни индикатори
        indirect = analysis.get("indirect_indicators", [])
        if indirect:
            self._heading(doc, "Индиректни индикатори", 2)
            for item in indirect:
                doc.add_paragraph(f"• {item.get('indicator', '')}", style="List Bullet")
                doc.add_paragraph(f"  {item.get('explanation', '')}")
            doc.add_paragraph()

        # Препоръки за преформулиране
        reforms = analysis.get("compliant_reformulations", [])
        if reforms:
            self._heading(doc, "Препоръчани формулировки", 2)
            tbl = doc.add_table(rows=len(reforms) + 1, cols=2)
            tbl.style = "Light Grid Accent 1"
            tbl.cell(0, 0).text = "Оригинално изискване"
            tbl.cell(0, 1).text = "Препоръчана формулировка"
            _set_cell_bg(tbl.cell(0, 0), "BDD7EE")
            _set_cell_bg(tbl.cell(0, 1), "BDD7EE")
            for i, r in enumerate(reforms, 1):
                tbl.cell(i, 0).text = r.get("original", "")
                tbl.cell(i, 1).text = r.get("compliant", "")
            doc.add_paragraph()

        doc.add_page_break()

    # ── Public API ────────────────────────────────────────────────────────────
    def generate(self, email_data: Dict, analysis_results: List[Dict]) -> str:
        doc = Document()

        # Стил на документа
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Метаданни
        props = doc.core_properties
        props.title = "Vendor-Lock Анализ – Обществени Поръчки"
        props.author = "Paraflow AI – Procurement Agent"
        props.comments = "Автоматично генериран анализ"

        self._add_title_page(doc, email_data, analysis_results)
        self._add_executive_summary(doc, analysis_results)

        for idx, result in enumerate(analysis_results):
            self._add_procurement_section(doc, result, idx)

        # Генерирай имена
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = f"vendor_lock_report_{ts}.docx"
        fpath = self.output_dir / fname
        doc.save(str(fpath))
        logger.info(f"Репорт записан: {fpath}")
        return str(fpath)
