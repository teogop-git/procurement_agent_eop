from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".xlsx"}


def normalize_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: List[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: List[str] = []

    for ws in wb.worksheets:
        parts.append(f"### SHEET: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_file_text(path: Path) -> str:
    suffix = path.suffix.lower()

    try:
        if suffix == ".txt":
            return extract_txt(path)
        if suffix == ".docx":
            return extract_docx(path)
        if suffix == ".pdf":
            return extract_pdf(path)
        if suffix == ".xlsx":
            return extract_xlsx(path)
    except Exception as exc:
        return f"[EXTRACT_ERROR] {path.name}: {exc}"

    return ""


def load_documents(documents_dir: Path) -> Tuple[str, Dict[str, str]]:
    file_texts: Dict[str, str] = {}
    merged_parts: List[str] = []

    for path in sorted(documents_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        text = normalize_text(extract_file_text(path))
        if not text:
            continue

        file_texts[path.name] = text
        merged_parts.append(f"\n\n===== SOURCE FILE: {path.name} =====\n{text}")

    return normalize_text("\n".join(merged_parts)), file_texts


def first_match(text: str, patterns: List[str], flags=re.IGNORECASE | re.DOTALL) -> str:
    for pattern in patterns:
        m = re.search(pattern, text, flags)
        if m:
            value = m.group(1).strip()
            value = re.sub(r"\s+\n", "\n", value)
            value = re.sub(r"\n\s+", "\n", value)
            value = re.sub(r"[ \t]{2,}", " ", value)
            return value.strip(" ;:-–—\n\t")
    return ""


def contains_any(text: str, terms: List[str]) -> bool:
    low = text.lower()
    return any(t.lower() in low for t in terms)


def compact(value: str, max_len: int = 2500) -> str:
    value = normalize_text(value)
    if len(value) > max_len:
        return value[:max_len].rstrip() + "..."
    return value


def section_after_label(text: str, labels: List[str], stop_labels: List[str], max_len: int = 2500) -> str:
    for label in labels:
        pattern = rf"{label}\s*[:\-–]?\s*(.+)"
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if not m:
            continue

        tail = m.group(1)
        stop_positions = []
        for stop in stop_labels:
            sm = re.search(rf"\n\s*{stop}\s*[:\-–]?", tail, re.IGNORECASE)
            if sm:
                stop_positions.append(sm.start())

        if stop_positions:
            tail = tail[: min(stop_positions)]

        return compact(tail, max_len=max_len)

    return ""


def extract_contacts(text: str) -> str:
    emails = sorted(set(re.findall(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)))
    phones = sorted(set(re.findall(r"(?:\+359|0)\s*\d[\d\s/().-]{6,}\d", text)))

    contact_name = first_match(text, [
        r"Лице за контакт\s*[:\-–]?\s*([^\n]+)",
        r"Контакт(?:но лице)?\s*[:\-–]?\s*([^\n]+)",
        r"Име\s*[:\-–]?\s*([А-ЯA-Z][^\n]{3,80})\s*\n\s*(?:Електронна поща|E-mail|Email)",
    ])

    parts = []
    if contact_name:
        parts.append(contact_name)
    if emails:
        parts.append("Електронна поща: " + ", ".join(emails[:5]))
    if phones:
        parts.append("Телефон: " + ", ".join(phones[:5]))

    return "\n".join(parts)


def extract_eop_reference(text: str, eop_url: str = "") -> str:
    nums = sorted(set(re.findall(r"\b\d{5}-\d{4}-\d{4}\b", text)))
    parts = []
    if nums:
        parts.append(nums[0])
    if eop_url:
        parts.append(eop_url)
    return "; ".join(parts)


def extract_ted_reference(text: str) -> str:
    ted_no = first_match(text, [
        r"\b(\d{6}-\d{4})\b",
        r"TED\s*[:\-–]?\s*([0-9]{6}-[0-9]{4})",
    ])

    oj = first_match(text, [
        r"(OJ\s*S\s*\d+/\d{4}\s*\d{2}/\d{2}/\d{4})",
        r"(ОВ\s*S\s*\d+/\d{4}\s*\d{2}/\d{2}/\d{4})",
    ])

    return " ; ".join([p for p in [ted_no, oj] if p])


def extract_buyer(text: str) -> str:
    return first_match(text, [
        r"Официално наименование\s*[:\-–]?\s*([^\n]+(?:\n[^\n]+){0,2})",
        r"Възложител\s*[:\-–]?\s*([^\n]+(?:\n[^\n]+){0,2})",
        r"Наименование на възложителя\s*[:\-–]?\s*([^\n]+(?:\n[^\n]+){0,2})",
    ])



def _clean_short_value(value: str, max_len: int = 1200) -> str:
    value = value or ""
    value = value.replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n\s+", "\n", value)
    value = re.sub(r"\s+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    value = value.strip(" \n\t;:-–—")

    junk_lines = [
        r"естество и количество",
        r"Когато основният предмет",
        r"Стратегическа поръчка",
        r"Поръчка с изисквания",
        r"Информация относно средства",
        r"Обществената поръчка е във връзка",
        r"Мотиви по чл\.\s*46",
    ]

    kept = []
    for line in value.splitlines():
        l = line.strip()
        if not l:
            continue
        if any(re.search(j, l, re.IGNORECASE) for j in junk_lines):
            continue
        if re.match(r"^[IVX]+\.\d+\)?$", l.strip(), re.IGNORECASE):
            continue
        kept.append(l)

    value = "\n".join(kept).strip()
    if len(value) > max_len:
        value = value[:max_len].rstrip() + "..."
    return value


def _one_line(value: str, max_len: int = 700) -> str:
    value = _clean_short_value(value, max_len=max_len)
    value = re.sub(r"\s*\n\s*", " ", value)
    value = re.sub(r"\s{2,}", " ", value)
    return value.strip(" ;:-–—")


def extract_subject(text: str) -> str:
    # 1) Най-надеждният източник: изречението "Предмет на настоящата обществена поръчка е ..."
    patterns = [
        r"Предмет\s+на\s+настоящата\s+обществена\s+поръчка\s+е\s+(.+?)(?:,\s*описан[аио]?\s+в|,\s*посочен[аио]?\s+в|\.\s|\n)",
        r"Предмет\s+на\s+обществената\s+поръчка\s+е\s+(.+?)(?:,\s*описан[аио]?\s+в|,\s*посочен[аио]?\s+в|\.\s|\n)",
    ]

    for pattern in patterns:
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if m:
            value = _one_line(m.group(1), max_len=800)
            # Предпазване от грешния match "артикулите, предмет на..."
            if len(value) > 20 and not value.lower().startswith(("посочени", "посочените", "описани")):
                return value

    # 2) Само line-anchored labels. Не търсим label-и по средата на изречение.
    line_patterns = [
        r"(?im)^\s*Предмет\s+на\s+поръчката\s*[:\-–]?\s*(.+)$",
        r"(?im)^\s*Предмет\s+на\s+обществената\s+поръчка\s*[:\-–]?\s*(.+)$",
        r"(?im)^\s*Наименование\s*[:\-–]?\s*(.+)$",
    ]

    for pattern in line_patterns:
        m = re.search(pattern, text)
        if m:
            value = _one_line(m.group(1), max_len=800)
            if len(value) > 20:
                return value

    # 3) Fallback: заглавие в кавички, но само ако изглежда като предмет.
    quoted = re.findall(r'[„"]([^„”"]{30,700})[”"]', text)
    for q in quoted:
        q1 = _one_line(q, max_len=800)
        if re.search(r"доставка|услуга|изграждане|поддръжка|консуматив|оборудване|система", q1, re.IGNORECASE):
            return q1

    return ""


def extract_short_description(text: str) -> str:
    # Цел: кратко описание, не цяла документация.
    # За този чеклист е най-полезно да се върне предмет + основни дейности.

    subject = extract_subject(text)

    activity_patterns = [
        r"Предметът\s+на\s+обществената\s+поръчка\s+включва\s+описание\s+на\s+дейностите\s+по\s+поръчката,?\s+както\s+следва\s*[:：]?\s*(.+?)(?=\n\s*(?:ОБЕКТ\s+НА|IV\.\d+\)?|Стратегическа\s+поръчка|Информация\s+относно|Разделяне\s+на\s+обособени|Мотиви\s+по\s+чл|Код\s+по\s+СРV|CPV|$))",
        r"Предметът\s+на\s+обществената\s+поръчка\s+включва\s*[:：]?\s*(.+?)(?=\n\s*(?:ОБЕКТ\s+НА|IV\.\d+\)?|Стратегическа\s+поръчка|Информация\s+относно|Разделяне\s+на\s+обособени|Мотиви\s+по\s+чл|Код\s+по\s+СРV|CPV|$))",
    ]

    activities = ""
    for pattern in activity_patterns:
        m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if m:
            activities = _clean_short_value(m.group(1), max_len=1800)
            break

    if activities:
        # Запазваме само реалните bullets а), б), в), г), ако ги има.
        bullet_matches = re.findall(
            r"([а-я]\)\s*.+?)(?=\n\s*[а-я]\)|\Z)",
            activities,
            flags=re.IGNORECASE | re.DOTALL,
        )

        if bullet_matches:
            bullets = [_clean_short_value(b, max_len=700) for b in bullet_matches]
            bullets = [b for b in bullets if b]
            if subject:
                return "Предмет: " + subject + "\n\nОсновни дейности:\n" + "\n".join(bullets)
            return "Основни дейности:\n" + "\n".join(bullets)

        if subject:
            return "Предмет: " + subject + "\n\n" + activities

        return activities

    # Fallback: кратък предмет + цел, но без раздели IV.*, финансиране, стратегическа поръчка, мотиви и т.н.
    desc = first_match(text, [
        r"Кратко\s+описание\s*[:\-–]?\s*(.+?)(?=\n\s*(?:IV\.\d+\)?|Стратегическа\s+поръчка|Информация\s+относно|Разделяне\s+на\s+обособени|Мотиви\s+по\s+чл|Вид\s+на\s+процедурата|Прогнозна\s+стойност|$))",
        r"Описание\s+на\s+предмета\s*[:\-–]?\s*(.+?)(?=\n\s*(?:IV\.\d+\)?|Стратегическа\s+поръчка|Информация\s+относно|Разделяне\s+на\s+обособени|Мотиви\s+по\s+чл|Вид\s+на\s+процедурата|Прогнозна\s+стойност|$))",
    ])

    desc = _clean_short_value(desc, max_len=1800)

    if subject and desc:
        if subject.lower() not in desc.lower():
            return "Предмет: " + subject + "\n\n" + desc
        return desc

    return subject


def extract_procedure_type(text: str) -> str:
    value = first_match(text, [
        r"Вид на процедурата\s*[:\-–]?\s*([^\n]+)",
        r"Начин на възлагане\s*[:\-–]?\s*([^\n]+)",
        r"Процедура\s*[:\-–]?\s*([^\n]+)",
    ])

    if value:
        return value

    if contains_any(text, ["открита процедура"]):
        return "Открита процедура"
    if contains_any(text, ["публично състезание"]):
        return "Публично състезание"
    if contains_any(text, ["събиране на оферти"]):
        return "Събиране на оферти с обява"

    return ""


def extract_estimated_value(text: str) -> str:
    value = first_match(text, [
        r"Прогнозна стойност[^\n:]*[:\-–]?\s*([0-9][0-9\s.,]+)\s*(евро|лв|bgn|eur)?",
        r"Стойност без ДДС[^\n:]*[:\-–]?\s*([0-9][0-9\s.,]+)\s*(евро|лв|bgn|eur)?",
    ])

    currency = first_match(text, [
        r"Прогнозна стойност[^\n:]*[:\-–]?\s*[0-9][0-9\s.,]+\s*(евро|лв|bgn|eur)",
        r"Стойност без ДДС[^\n:]*[:\-–]?\s*[0-9][0-9\s.,]+\s*(евро|лв|bgn|eur)",
    ])

    if value and currency:
        return f"{value} {currency} без ДДС"
    return value


def extract_contract_duration(text: str) -> str:
    return first_match(text, [
        r"Продължителност[^\n:]*[:\-–]?\s*([0-9]+\s*(?:месец|месеца|дни|години)[^\n]*)",
        r"Срок на договора[^\n:]*[:\-–]?\s*([0-9]+\s*(?:месец|месеца|дни|години)[^\n]*)",
        r"срок от\s*([0-9]+\s*(?:месец|месеца|дни|години))",
    ])


def extract_dates(text: str) -> Dict[str, str]:
    result = {}

    result["offer_deadline"] = first_match(text, [
        r"Краен срок за получаване на оферти\s*[:\-–]?\s*([^\n]+)",
        r"Срок за получаване на оферти\s*[:\-–]?\s*([^\n]+)",
        r"Краен срок за подаване на оферти\s*[:\-–]?\s*([^\n]+)",
    ])

    result["opening_date"] = first_match(text, [
        r"Отваряне на офертите\s*[:\-–]?\s*([^\n]+)",
        r"Дата и час на отваряне\s*[:\-–]?\s*([^\n]+)",
    ])

    result["publication_date"] = first_match(text, [
        r"Дата на публикуване\s*[:\-–]?\s*([^\n]+)",
        r"Публикувано на\s*[:\-–]?\s*([^\n]+)",
    ])

    result["questions_deadline"] = first_match(text, [
        r"Срок за въпроси\s*[:\-–]?\s*([^\n]+)",
        r"Краен срок за искане на разяснения\s*[:\-–]?\s*([^\n]+)",
    ])

    result["change_deadline"] = first_match(text, [
        r"Срок за промяна\s*[:\-–]?\s*([^\n]+)",
    ])

    return result


def extract_delivery_term(text: str) -> str:
    return compact(first_match(text, [
        r"Срок за изпълнение на доставките\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Място|Гаранция|Плащане|Срок за плащане|При|В случай|$))",
        r"Срок за доставка\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Място|Гаранция|Плащане|Срок за плащане|При|В случай|$))",
        r"достав[а-я]+\s+.*?в срок\s+до\s+([0-9]+[^.\n]+)",
    ]), 1200)


def extract_delivery_place(text: str) -> str:
    return compact(first_match(text, [
        r"Място на доставка[^\n:]*[:\-–]?\s*(.+?)(?:\n\s*(?:Срок|Гаранция|Плащане|Цена|$))",
        r"Място на изпълнение[^\n:]*[:\-–]?\s*(.+?)(?:\n\s*(?:Срок|Гаранция|Плащане|Цена|$))",
    ]), 1500)


def extract_guarantees(text: str) -> Dict[str, str]:
    result = {}

    result["performance_guarantee"] = compact(first_match(text, [
        r"Гаранцията за изпълнение на договора\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Финансиране|Разглеждане|Критерий|$))",
        r"Гаранция за изпълнение\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Финансиране|Разглеждане|Критерий|$))",
        r"гаранция[^\n.]{0,80}изпълнение[^\n.]{0,80}в размер на\s*([^.\n]+)",
    ]), 1200)

    result["good_performance_guarantee"] = compact(first_match(text, [
        r"Гаранция за добро изпълнение\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Срок на договора|Неустойки|Други|$))",
        r"Гаранция за обезпечаване изпълнението\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Срок на договора|Неустойки|Други|$))",
    ]), 1200)

    result["advance_payment_guarantee"] = compact(first_match(text, [
        r"Гаранция за авансово плащане\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Гаранция за добро|Срок на договора|Неустойки|$))",
    ]), 1000)

    return result


def extract_award(text: str) -> Dict[str, str]:
    result = {}

    if contains_any(text, ["най-ниска цена", "най-ниската цена"]):
        result["award_price"] = "Най-ниска цена"
    else:
        result["award_price"] = first_match(text, [
            r"Критерий за възлагане\s*[:\-–]?\s*([^\n]+)",
            r"Критерии за възлагане\s*[:\-–]?\s*([^\n]+)",
        ])

    result["evaluation_methodology"] = compact(first_match(text, [
        r"Методика\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Критерии за подбор|Административни|Икономическо|Технически и професионални|$))",
        r"П\s*=\s*(.+?)(?:\n\s*(?:Критерии за подбор|Административни|Икономическо|$))",
    ]), 2500)

    return result


def extract_selection_criteria(text: str) -> Dict[str, str]:
    result = {}

    result["article_5k_declaration"] = "ДА" if contains_any(text, [
        "чл. 5к",
        "чл.5к",
        "Регламент (ЕС) № 833/2014",
        "Регламент ЕС № 833/2014",
    ]) else ""

    result["specific_turnover"] = compact(first_match(text, [
        r"(?:Специфичен оборот|конкретен годишен оборот|общ оборот[^\n]{0,80}сферата)[^\n:]*[:\-–]?\s*(.+?)(?:\n\s*(?:Технически и професионални|Участие|Изпълнени дейности|Екип|$))",
        r"минимален общ оборот\s*(.+?)(?:\n\s*(?:Технически и професионални|Участие|Изпълнени дейности|Екип|$))",
    ]), 2000)

    result["similar_deliveries"] = compact(first_match(text, [
        r"Изпълнени дейности/услуги\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Екип|Оторизации|Инструменти|Сертификат|$))",
        r"през последните\s*3\s*\(?.{0,20}години.{0,2000}?(?:доставка|продажба).+?(?=\n\s*(?:Екип|Оторизации|Инструменти|Сертификат|$))",
    ]), 2500)

    result["authorizations"] = compact(first_match(text, [
        r"Оторизации\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Инструменти|Сертификат|ТЕХНИЧЕСКА|$))",
        r"(документ[^\n]{0,200}оторизирани[^\n]+(?:\n[^\n]+){0,5})",
    ]), 1800)

    return result


def extract_technical(text: str) -> Dict[str, str]:
    result = {}

    result["brand_model_manufacturer"] = first_match(text, [
        r"Марка, модел, производител\s*[:\-–]?\s*([^\n]+)",
        r"с марки\s*([^\n.,]+)",
    ])

    result["technical_requirements"] = compact(first_match(text, [
        r"Технически изисквания\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Вендор|Срок за реакция|Гаранционен|Срок за доставка|Сравнителна|$))",
        r"Под\s+„?оригинални“?.+?(?=\n\s*(?:Вендор|Срок за реакция|Гаранционен|Срок за доставка|Сравнителна|$))",
        r"Специални изисквания към предлаганите.+?(?=\n\s*(?:Вендор|Срок за реакция|Гаранционен|Срок за доставка|Сравнителна|$))",
    ]), 3500)

    result["reaction_term"] = compact(first_match(text, [
        r"Срок за реакция и условия\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Гаранционен|Срок за доставка|Сравнителна|$))",
        r"констатирани липси.+?срок до\s*([0-9]+[^\n.]+)",
    ]), 1500)

    result["authorization_timing"] = "При подаване на оферта" if contains_any(text, [
        "При подаване на оферта",
        "при подаване на оферта",
    ]) else ""

    return result


def extract_contract(text: str) -> Dict[str, str]:
    result = {}

    result["payment_term"] = compact(first_match(text, [
        r"Срок за плащане\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Авансово|Гаранция за авансово|Гаранция за добро|Срок на договора|$))",
        r"плащането.+?в срок до\s*([0-9]+[^\n.]+)",
    ]), 1200)

    result["advance_payment"] = compact(first_match(text, [
        r"Авансово плащане\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Гаранция за авансово|Гаранция за добро|Срок на договора|$))",
    ]), 1000)

    result["contract_term"] = compact(first_match(text, [
        r"Срок на договора\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Неустойки|Други специфични|ДРУГИ|$))",
        r"срок от\s*([0-9]+\s*(?:месец|месеца|дни|години).+?)(?:\.|\n)",
    ]), 1500)

    result["penalties"] = compact(first_match(text, [
        r"Неустойки\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Други специфични|ДРУГИ|ЕЗИК|$))",
        r"неустойка в размер на\s*([^.\n]+(?:\n[^\n]+){0,5})",
    ]), 1800)

    result["language_requirements"] = compact(first_match(text, [
        r"ЕЗИК\s*[:\-–]?\s*(.+?)(?:\n\s*(?:$|ДРУГИ|Приложения))",
        r"Офертите се изготвят на български език.+?(?:\.|\n)",
    ]), 1000)

    return result


def extract_price(text: str) -> Dict[str, str]:
    result = {}

    result["price_offer_template"] = "ДА" if contains_any(text, [
        "образец ценово предложение",
        "ценово предложение",
        "Приложение 2",
        "Приложение № 2",
    ]) else ""

    result["price_offer_attachments"] = compact(first_match(text, [
        r"ПРИЛОЖЕНИЯ към ценовото предложение\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Указания|Специфични|ПРОЕКТО-ДОГОВОР|$))",
    ]), 1200)

    result["offer_submission_instructions"] = compact(first_match(text, [
        r"Указания за изготвяне и подаване на офертите.+?\s*(.+?)(?:\n\s*(?:Специфични изисквания|ПРОЕКТО-ДОГОВОР|$))",
        r"Цената включва\s*(.+?)(?:\n\s*(?:Специфични изисквания|ПРОЕКТО-ДОГОВОР|$))",
    ]), 1800)

    return result



def extract_lots(text: str) -> str:
    # Това поле НЕ трябва да съдържа мотиви за неразделяне.
    # Само: списък с позиции или "Няма обособени позиции."

    lot_items = []

    # Формати: "Обособена позиция № 1: ...", "ОП 1 - ..."
    patterns = [
        r"Обособена\s+позиция\s*№?\s*([0-9]+)\s*[:\.\-–]?\s*(.+?)(?=\n\s*Обособена\s+позиция\s*№?\s*[0-9]+|\n\s*ОП\s*№?\s*[0-9]+|\n\s*(?:III\.|IV\.|Раздел|Критерий|Прогнозна|$))",
        r"ОП\s*№?\s*([0-9]+)\s*[:\.\-–]?\s*(.+?)(?=\n\s*ОП\s*№?\s*[0-9]+|\n\s*Обособена\s+позиция\s*№?\s*[0-9]+|\n\s*(?:III\.|IV\.|Раздел|Критерий|Прогнозна|$))",
    ]

    for pattern in patterns:
        for num, title in re.findall(pattern, text, flags=re.IGNORECASE | re.DOTALL):
            title = _one_line(title, max_len=300)
            title = re.sub(r"^(Наименование|Предмет)\s*[:\-–]?\s*", "", title, flags=re.IGNORECASE).strip()
            if title and not re.search(r"мотиви|нецелесъобразно|неразделянето", title, re.IGNORECASE):
                lot_items.append((int(num), title))

    if lot_items:
        seen = set()
        lines = []
        for num, title in sorted(lot_items, key=lambda x: x[0]):
            key = (num, title.lower())
            if key in seen:
                continue
            seen.add(key)
            lines.append(f"{num}. {title}")
        return "\n".join(lines)

    # Ако документацията казва, че няма разделяне/позиции.
    no_lots_terms = [
        "не се предвижда разделяне",
        "не се разделя на обособени позиции",
        "няма обособени позиции",
        "разделянето на обществената поръчка на обособени позиции е нецелесъобразно",
        "мотиви за невъзможността за разделяне",
        "неразделянето на обществената поръчка",
    ]

    low = text.lower()
    if any(term in low for term in no_lots_terms):
        return "Няма обособени позиции."

    # Line-anchored fallback, но пак без мотиви.
    value = first_match(text, [
        r"(?im)^\s*Обособени\s+позиции\s*[:\-–]?\s*(.+)$",
        r"(?im)^\s*Разделяне\s+на\s+обособени\s+позиции\s*[:\-–]?\s*(.+)$",
    ])

    value = _one_line(value, max_len=500)

    if not value:
        return ""

    if re.search(r"не|няма|нецелесъобразно|мотиви", value, re.IGNORECASE):
        return "Няма обособени позиции."

    return value


def build_checklist_data(text: str, eop_url: str = "") -> Dict[str, str]:
    data: Dict[str, str] = {}

    data.update({
        "sales_owner": "",
        "coordinator": "",
        "presales_consultant": "",
        "product_manager": "",
        "technical_specialist": "",

        "buyer": extract_buyer(text),
        "contacts": extract_contacts(text),
        "eop_reference": extract_eop_reference(text, eop_url=eop_url),
        "ted_reference": extract_ted_reference(text),
        "subject": extract_subject(text),
        "lots": extract_lots(text),
        "short_description": extract_short_description(text),
        "procedure_type": extract_procedure_type(text),
        "estimated_value": extract_estimated_value(text),
        "payform_required": "",
        "contract_duration": extract_contract_duration(text),
        "execution_term": extract_execution_term(text),
        "delivery_term": extract_delivery_term(text),
        "offer_validity": extract_offer_validity(text),
        "funding": first_match(text, [
            r"Финансиране\s*[:\-–]?\s*([^\n]+(?:\n[^\n]+){0,3})",
        ]),
        "offer_review": "",
        "technical_score": "",

        "third_party_capacity": "",
        "subcontracting": "",
        "consortium": "",
        "third_party_capacity_alt": "",
        "total_turnover": "",
        "market_consultation": "",
        "team_qualification": "",
        "tools_equipment": "",
        "iso_certificate": "",

        "vendor": "",
        "warranty_terms": "",
        "delivery_installation_term": "",
        "comparison_table": "",
        "product_certificates": "",
        "delivery_execution_place": extract_delivery_place(text),
        "brochures": "",
        "helpdesk": "",
        "technical_project": "",
        "project_schedule": "",
        "execution_methodology": "",
        "risk_management_methodology": "",
        "quality_management_methodology": "",
        "confidentiality_declaration": "",
        "tax_social_security_declaration": "",

        "draft_contract_specific_requirements": "",
        "other_specific_requirements": "",
        "contract_delivery_support_term": "",
        "delivery_place": extract_delivery_place(text),
    })

    data.update(extract_dates(text))
    data.update(extract_guarantees(text))
    data.update(extract_award(text))
    data.update(extract_selection_criteria(text))
    data.update(extract_technical(text))
    data.update(extract_contract(text))
    data.update(extract_price(text))

    if data.get("estimated_value"):
        m = re.search(r"([0-9][0-9\s.,]+)", data["estimated_value"])
        if m:
            number = m.group(1).replace(" ", "").replace(",", ".")
            try:
                if float(number) > 50000:
                    data["payform_required"] = "ДА"
            except ValueError:
                pass

    return data



# --- BEGIN CHECKLIST STRICT FIELD OVERRIDES ---

_BG_DATE_RE = (
    r"\d{1,2}\s*"
    r"(?:януари|февруари|март|април|май|юни|юли|август|септември|октомври|ноември|декември)"
    r"\s+\d{4}"
    r"(?:\s*\([^)]+\))?"
    r"(?:,\s*\d{1,2}:\d{2})?"
)

_NUM_DATE_RE = r"\d{1,2}[./-]\d{1,2}[./-]\d{4}(?:,\s*\d{1,2}:\d{2})?"

_ANY_DATE_RE = rf"(?:{_BG_DATE_RE}|{_NUM_DATE_RE})"


def _strip_bt_noise(value: str) -> str:
    value = value or ""
    value = re.sub(r"\(BT-[^)]+\)", " ", value)
    value = re.sub(r"\bBT-[A-Za-z0-9()_.-]+", " ", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip(" ;:-–—")


def _first_date(value: str) -> str:
    value = _strip_bt_noise(value)
    m = re.search(_ANY_DATE_RE, value, re.IGNORECASE)
    return m.group(0).strip() if m else ""


def _find_date_near_label(text: str, labels: list[str], window: int = 900) -> str:
    cleaned = _strip_bt_noise(text)

    for label in labels:
        pattern = re.escape(label) + rf"\s*[:\-–]?\s*(.{{0,{window}}})"
        m = re.search(pattern, cleaned, re.IGNORECASE | re.DOTALL)
        if not m:
            continue

        segment = m.group(1)
        date_value = _first_date(segment)
        if date_value:
            return date_value

    return ""


def _find_duration_near_label(text: str, labels: list[str], window: int = 500) -> str:
    cleaned = _strip_bt_noise(text)

    for label in labels:
        pattern = re.escape(label) + rf"\s*[:\-–]?\s*(.{{0,{window}}})"
        m = re.search(pattern, cleaned, re.IGNORECASE | re.DOTALL)
        if not m:
            continue

        segment = m.group(1)
        dm = re.search(
            r"\b\d+\s*(?:месец|месеца|дни|ден|година|години)\b",
            segment,
            re.IGNORECASE,
        )
        if dm:
            return dm.group(0).strip()

    return ""


def extract_execution_term(text: str) -> str:
    # За "Срок за изпълнение" в чеклиста искаме начална дата, ако е налична.
    date_value = _find_date_near_label(text, ["Начална дата"], window=250)
    if date_value:
        return f"Начална дата {date_value}"

    date_value = _find_date_near_label(text, ["Срок за изпълнение"], window=500)
    if date_value:
        return date_value

    return ""


def extract_offer_validity(text: str) -> str:
    value = _find_duration_near_label(
        text,
        ["Валидност на офертата", "Срок на валидност на офертата"],
        window=400,
    )

    if value:
        # Нормализиране само за визуално съвпадение с шаблона.
        value = value.replace("месеца", "Месеца")
        value = value.replace("месец", "Месец")
        return value

    return ""


def extract_dates(text: str) -> Dict[str, str]:
    return {
        "offer_deadline": _find_date_near_label(
            text,
            [
                "Краен срок за получаване на оферти",
                "Краен срок за подаване на оферти",
                "Срок за получаване на оферти",
            ],
            window=900,
        ),
        "opening_date": _find_date_near_label(
            text,
            [
                "Отваряне на офертите",
                "Дата и час на отваряне",
                "Дата на отваряне на офертите",
            ],
            window=900,
        ),
        "publication_date": _find_date_near_label(
            text,
            [
                "Дата на публикуване",
                "Публикувано на",
            ],
            window=500,
        ),
        "questions_deadline": _find_date_near_label(
            text,
            [
                "Срок за въпроси",
                "Краен срок за искане на разяснения",
            ],
            window=500,
        ),
        "change_deadline": _find_date_near_label(
            text,
            [
                "Срок за промяна",
            ],
            window=500,
        ),
    }


def _guarantee_sentence_near_label(text: str, labels: list[str], window: int = 1200) -> str:
    cleaned = _strip_bt_noise(text)

    for label in labels:
        pattern = re.escape(label) + rf"\s*[:\-–]?\s*(.{{0,{window}}})"
        m = re.search(pattern, cleaned, re.IGNORECASE | re.DOTALL)
        if not m:
            continue

        segment = m.group(1)

        gm = re.search(
            r"(в\s+размер\s+на\s+\d+(?:[,.]\d+)?\s*%[^.\n]{0,350})",
            segment,
            re.IGNORECASE | re.DOTALL,
        )
        if gm:
            value = gm.group(1)
            value = re.sub(r"\s+", " ", value).strip(" ;:-–—")
            if not value.endswith("."):
                value += "."
            return value

        gm = re.search(
            r"(?:гаранция[^\n.]{0,120}изпълнение[^\n.]{0,120})(\d+(?:[,.]\d+)?\s*%[^.\n]{0,350})",
            segment,
            re.IGNORECASE | re.DOTALL,
        )
        if gm:
            value = "в размер на " + gm.group(1)
            value = re.sub(r"\s+", " ", value).strip(" ;:-–—")
            if not value.endswith("."):
                value += "."
            return value

    return ""


def extract_guarantees(text: str) -> Dict[str, str]:
    performance = _guarantee_sentence_near_label(
        text,
        [
            "Гаранцията за изпълнение на договора",
            "Гаранция за изпълнение на договора",
            "Гаранция за изпълнение",
        ],
    )

    good_performance = _guarantee_sentence_near_label(
        text,
        [
            "Гаранция за добро изпълнение",
            "Гаранция за обезпечаване изпълнението",
            "Гаранция за обезпечаване на изпълнението",
        ],
    )

    advance = first_match(text, [
        r"Гаранция за авансово плащане\s*[:\-–]?\s*(.+?)(?:\n\s*(?:Гаранция за добро|Срок на договора|Неустойки|$))",
    ])

    return {
        "performance_guarantee": performance,
        "good_performance_guarantee": good_performance or performance,
        "advance_payment_guarantee": compact(advance, 1000),
    }

# --- END CHECKLIST STRICT FIELD OVERRIDES ---


def write_preview(data: Dict[str, str]) -> str:
    lines = []
    for k in sorted(data.keys()):
        v = data.get(k, "")
        if isinstance(v, str):
            vv = v.replace("\n", " ")
            if len(vv) > 160:
                vv = vv[:160] + "..."
        else:
            vv = str(v)
        status = "OK" if v else "EMPTY"
        lines.append(f"{status:6} | {k:35} | {vv}")
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract real procurement checklist data from local OP documents.")
    parser.add_argument("--documents-dir", required=True, help="Directory containing real EOP documents: PDF/DOCX/XLSX/TXT")
    parser.add_argument("--eop-url", default="", help="EOP URL, e.g. https://app.eop.bg/today/557865")
    parser.add_argument("--output-json", required=True, help="Output checklist_data JSON")
    parser.add_argument("--debug-text", default="", help="Optional path to write merged extracted text")
    parser.add_argument("--preview", action="store_true", help="Print extracted key/value preview")
    args = parser.parse_args()

    documents_dir = Path(args.documents_dir)
    if not documents_dir.exists():
        raise SystemExit(f"documents-dir does not exist: {documents_dir}")

    merged_text, file_texts = load_documents(documents_dir)
    if not merged_text:
        raise SystemExit(f"No extractable text found in: {documents_dir}")

    data = build_checklist_data(merged_text, eop_url=args.eop_url)

    output_json = Path(args.output_json)
    output_json.parent.mkdir(parents=True, exist_ok=True)
    output_json.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    if args.debug_text:
        debug_path = Path(args.debug_text)
        debug_path.parent.mkdir(parents=True, exist_ok=True)
        debug_path.write_text(merged_text, encoding="utf-8")

    result = {
        "documents_dir": str(documents_dir),
        "files_read": list(file_texts.keys()),
        "output_json": str(output_json),
        "filled_values": sum(1 for v in data.values() if isinstance(v, str) and v.strip()),
        "empty_values": sum(1 for v in data.values() if not (isinstance(v, str) and v.strip())),
    }

    print(json.dumps(result, ensure_ascii=False, indent=2))

    if args.preview:
        print("\n--- PREVIEW ---")
        print(write_preview(data))

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
